import os.path
import docx
# from win32com import client



def docx2string(path):
    """
    提取docx文件的文本并转化为字符串

    :param path：docx文件的完整路径
    :return: 提取的整个文本的字符串
    """
    with open(path, 'rb') as file:
        doc = docx.Document(file)
        result = ''
        # 读取段落
        for p in doc.paragraphs:
            text = p.text
            result += text
        # 对于一些文档还要读取表格
        for tab in doc.tables:
            for row in tab.rows:
                for cell in row.cells:
                    text = p.text
                    result += text
        return result


# def doc2docx(path):
#     """
#     将doc文件转换为为同名docx文件
#
#     :param path: doc文件的完整路径
#     :raise FileNotFoundError
#     """
#     if not os.path.exists(path):
#         raise FileNotFoundError
#     else:
#         word = client.DispatchEx("Word.Application")
#         doc = word.Documents.Open(path)
#         doc.SaveAs(path[:-3] + 'docx', 16)
#         doc.Close()
#         word.Quit()


def is_already_converted(path):
    """
    检测doc文件是否已经转换

    :param path: doc文件的完整路径
    :return: 是否已经转换，已经转换返回True
    """
    if os.path.exists(path[:-3] + "docx"):
        return True
    else:
        return False

